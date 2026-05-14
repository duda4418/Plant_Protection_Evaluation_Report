from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate, InlineImage
from jinja2 import Environment, StrictUndefined


class RenderError(RuntimeError):
    pass


def render_report(context: dict, template_path: Path, output_path: Path) -> None:
    if not template_path.exists():
        raise RenderError(f"Template not found: {template_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = DocxTemplate(template_path)
    render_context = deepcopy(context)

    for product in render_context["products"]:
        chart_path = product.get("chart_path")
        if chart_path:
            product["efficacy_chart"] = InlineImage(doc, chart_path, width=docx_mm(145))

    jinja_env = Environment(undefined=StrictUndefined, autoescape=False)
    doc.render(render_context, jinja_env=jinja_env)
    doc.save(output_path)
    _assert_no_unresolved_placeholders(output_path)


def docx_mm(value: int):
    from docx.shared import Mm

    return Mm(value)


def _assert_no_unresolved_placeholders(output_path: Path) -> None:
    text = _extract_docx_text(output_path)
    unresolved_markers = ("{{", "}}", "{%", "%}")
    found = [marker for marker in unresolved_markers if marker in text]
    if found:
        markers = ", ".join(found)
        raise RenderError(f"Generated DOCX still contains unresolved template markers: {markers}")


def _extract_docx_text(path: Path) -> str:
    document = Document(path)
    chunks: list[str] = []
    chunks.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in document.sections:
        chunks.extend(paragraph.text for paragraph in section.header.paragraphs)
        chunks.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(chunks)
