from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .template_styles import (
    BRAND_NAVY,
    BRAND_ORANGE,
    DARK_RED,
    TEXT_DARK,
    TEXT_MUTED,
    configure_document_theme,
    set_paragraph_bottom_border,
    style_cell_text,
    style_run,
    style_metric_table,
    style_summary_table,
    style_table_header,
    style_warning_cell,
)


def build_default_template(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_document_theme(document)
    _add_header_footer(document)
    _add_title_page(document)
    _add_products_section(document)
    document.save(output_path)


# Document shell
def _add_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.clear()
    header_run = header.add_run("{{ metadata.issuing_authority }}")
    style_run(header_run, color=TEXT_MUTED, size=7)
    separator_run = header.add_run(" | {{ metadata.report_id }}")
    style_run(separator_run, color=TEXT_MUTED, size=8)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer.paragraphs[0]
    footer.clear()
    footer_table = section.footer.add_table(rows=1, cols=2, width=Inches(6.3))
    left_cell = footer_table.rows[0].cells[0]
    right_cell = footer_table.rows[0].cells[1]
    left_cell.text = "{{ footer.classification }}"
    for run in left_cell.paragraphs[0].runs:
        style_run(run, color=TEXT_MUTED, size=7)
    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_text = right_paragraph.add_run("Page ")
    style_run(footer_text, color=TEXT_MUTED, size=7)
    _add_field(right_paragraph, "PAGE")
    page_text = right_paragraph.add_run(" / ")
    style_run(page_text, color=TEXT_MUTED, size=7)
    _add_field(right_paragraph, "NUMPAGES")


def _add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


# Front matter
def _add_title_page(document: Document) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(28)
    spacer.paragraph_format.space_after = Pt(0)

    title = document.add_paragraph(style="Title")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title.add_run("{{ metadata.report_title }}")

    metadata_line = document.add_paragraph()
    metadata_line.paragraph_format.space_after = Pt(14)
    metadata_line_run = metadata_line.add_run(
        "Issuing authority: {{ metadata.issuing_authority }} · Report ID: {{ metadata.report_id }} · Issued on: {{ metadata.issued_on }} (v{{ metadata.version }})"
    )
    style_run(metadata_line_run, color=TEXT_MUTED, size=7)

    summary_box = document.add_table(rows=1, cols=2)
    summary_box.style = "Table Grid"
    summary_box.rows[0].cells[0].text = ""
    summary_cell = summary_box.rows[0].cells[1]
    summary_cell.paragraphs[0].clear()
    summary_bold_run = summary_cell.paragraphs[0].add_run("Overall summary.")
    style_run(summary_bold_run, color=TEXT_DARK, size=8, bold=True)
    summary_text_run = summary_cell.paragraphs[0].add_run(" {{ summary.summary_text }}")
    style_run(summary_text_run, color=TEXT_DARK, size=8)
    style_summary_table(summary_box)

    document.add_page_break()


# Repeating product section
def _add_products_section(document: Document) -> None:
    document.add_paragraph("{%p for product in products %}")
    document.add_paragraph("{%p if not loop.first %}")
    document.add_page_break()
    document.add_paragraph("{%p endif %}")
    heading = document.add_heading("{{ product.name }}", level=1)
    set_paragraph_bottom_border(heading, "D9D9D9")
    _add_product_overview(document)
    _add_high_risk_notice(document)
    _add_product_data_sections(document)
    document.add_paragraph("{%p if product.show_high_risk_notice %}")
    document.add_page_break()
    document.add_paragraph("{%p endif %}")
    _add_chart_section(document)
    _add_restrictions_section(document)
    document.add_paragraph("{%p if not product.show_restrictions %}")
    document.add_page_break()
    document.add_paragraph("{%p endif %}")
    _add_environmental_notes_section(document)
    document.add_paragraph("{%p endfor %}")


# Product overview
def _add_product_overview(document: Document) -> None:
    facts = document.add_paragraph()
    facts.paragraph_format.space_after = Pt(2)
    facts.add_run("Active substance: ")
    active_sub_run = facts.add_run("{{ product.active_substance }}")
    style_run(active_sub_run, color=TEXT_DARK, size=8, bold=True)
    facts.add_run(" ({{ product.concentration }}) · Formulation: {{ product.formulation_display }} · Manufacturer: {{ product.manufacturer }}")

    approval = document.add_paragraph()
    approval.paragraph_format.space_after = Pt(8)
    approval_run = approval.add_run("{{ product.approval_summary }}")
    style_run(approval_run, color=BRAND_NAVY, size=8)
    approval_run.italic = True

    registration = document.add_paragraph()
    registration.paragraph_format.space_after = Pt(12)
    reg_prefix_run = registration.add_run("{{ product.registration_prefix }}")
    style_run(reg_prefix_run, color=TEXT_DARK, size=8)
    reg_bold_run = registration.add_run("{{ product.average_efficacy }}")
    style_run(reg_bold_run, color=TEXT_DARK, size=8, bold=True)
    reg_suffix_run = registration.add_run(".")
    style_run(reg_suffix_run, color=TEXT_DARK, size=8)


def _add_high_risk_notice(document: Document) -> None:
    document.add_paragraph("{%p if product.show_high_risk_notice %}")
    notice_table = document.add_table(rows=1, cols=1)
    notice_table.style = "Table Grid"
    notice_cell = notice_table.rows[0].cells[0]
    notice_cell.text = ""
    notice_heading = notice_cell.paragraphs[0].add_run("HIGH RISK NOTICE")
    style_run(notice_heading, color=DARK_RED, size=8, bold=True)
    notice_body = notice_cell.add_paragraph().add_run("{{ product.high_risk_notice }}")
    style_run(notice_body, color="1A1A1A", size=8)
    style_warning_cell(notice_cell)
    document.add_paragraph("{%p endif %}")


# Product data blocks
def _add_product_data_sections(document: Document) -> None:
    document.add_heading("Applications", level=2)
    _add_applications_table(document)

    document.add_heading("Toxicity Indicators", level=2)
    _add_toxicity_table(document)

    document.add_heading("Risk Components", level=2)
    _add_risk_table(document)


def _add_chart_section(document: Document) -> None:
    document.add_heading("Efficacy by crop", level=2)
    chart_paragraph = document.add_paragraph()
    chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart_paragraph.add_run("{{ product.efficacy_chart }}")
    chart_note = document.add_paragraph("{{ product.chart_note }}")
    chart_note.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_restrictions_section(document: Document) -> None:
    document.add_paragraph("{%p if product.show_restrictions %}")
    document.add_heading("{{ product.restrictions_heading }}", level=2)
    document.add_paragraph("{%p for restriction in product.restrictions %}")
    document.add_paragraph("• {{ restriction }}")
    document.add_paragraph("{%p endfor %}")
    document.add_paragraph("{%p endif %}")


def _add_environmental_notes_section(document: Document) -> None:
    document.add_heading("Environmental notes", level=2)
    document.add_paragraph("{%p for note in product.environmental_notes %}")
    document.add_paragraph("• {{ note }}")
    document.add_paragraph("{%p endfor %}")


# Reusable tables
def _add_applications_table(document: Document) -> None:
    table = document.add_table(rows=4, cols=7)
    table.style = "Table Grid"
    headers = ["Crop", "Target pest", "Stage\n(BBCH)", "Dose", "Max/\nseason", "Efficacy", "PHI\n(days)"]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header
    style_table_header(table.rows[0], fill=BRAND_NAVY)
    table.rows[1].cells[0].text = "{%tr for app in product.applications %}"
    values = [
        "{{ app.crop }}",
        "{{ app.target_pest }}",
        "{{ app.bbch }}",
        "{{ app.dose }}",
        "{{ app.max_applications }}",
        "{{ app.efficacy }}",
        "{{ app.phi_days }}",
    ]
    for cell, value in zip(table.rows[2].cells, values, strict=True):
        cell.text = value
    table.rows[3].cells[0].text = "{%tr endfor %}"
    style_metric_table(table)
    # Make target_pest column italic in the data row (applied after style_metric_table to avoid override)
    style_cell_text(table.rows[2].cells[1], color=TEXT_DARK, bold=False, size=7, italic=True)


def _add_risk_table(document: Document) -> None:
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "{%tr for row in product.risk_component_rows %}"
    table.rows[1].cells[0].text = "{{ row.component }}"
    table.rows[1].cells[1].text = "{{ row.score }}"
    table.rows[2].cells[0].text = "{%tr endfor %}"
    style_metric_table(table)


def _add_toxicity_table(document: Document) -> None:
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "{%tr for row in product.toxicity_rows %}"
    table.rows[1].cells[0].text = "{{ row.indicator }}"
    table.rows[1].cells[1].text = "{{ row.value }}"
    table.rows[2].cells[0].text = "{%tr endfor %}"
    style_metric_table(table)
